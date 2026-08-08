import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for b_name, b_val in borders.items():
        if b_val is not None:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), b_val.get('val', 'single'))
            node.set(qn('w:sz'), str(b_val.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), b_val.get('color', 'auto'))
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_body_p(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    r.bold = bold
    r.italic = italic
    return p

def add_bullet_p(doc, prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return p

def add_callout_box(doc, title, text, bg_hex="F8FAFC", border_hex="1B365D", icon="💡"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    set_cell_borders(cell, left={'val': 'single', 'sz': 24, 'color': border_hex})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{icon} {title}\n")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor.from_string(border_hex)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_styled_table(doc, headers, rows_data, col_widths=None):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], '1B365D')
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row in enumerate(rows_data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg_color = 'F8FAFC' if r_idx % 2 == 1 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            set_cell_borders(row_cells[c_idx], 
                             bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'})
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_image_figure(doc, img_path, caption_text, width_inches=6.0):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(caption_text)
        r_cap.italic = True
        r_cap.font.name = 'Arial'
        r_cap.font.size = Pt(8.5)
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc = docx.Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ---------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------
if os.path.exists('static/LogoAP.png'):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.add_run()
    run_logo.add_picture('static/LogoAP.png', width=Inches(1.8))

p_title_space = doc.add_paragraph()
p_title_space.paragraph_format.space_before = Pt(40)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_after = Pt(6)
r_title = p_title.add_run('AN PHƯỚC RETAIL COMMANDER')
r_title.bold = True
r_title.font.name = 'Arial'
r_title.font.size = Pt(26)
r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_after = Pt(20)
r_sub = p_sub.add_run('HƯỚNG DẪN SỬ DỤNG WEB PORTAL NHẬP LIỆU & BÁO CÁO VẬN HÀNH QLKD')
r_sub.bold = True
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Emerald

# Divider bar
p_div = doc.add_paragraph()
p_div.paragraph_format.space_after = Pt(40)
r_div = p_div.add_run('────────────────────────────────────────────────────────')
r_div.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

# Meta Table Box
meta_headers = ['Thông Tin Hệ Thống', 'Giá Trị Chi Tiết']
meta_rows = [
    ['Phiên bản hệ thống', 'Retail Commander Web Portal v2026.8 (Cập nhật 08/2026)'],
    ['Môi trường vận hành', 'Production Cloud Server (Render & Neon PostgreSQL)'],
    ['Đơn vị phát hành', 'Phòng Quản Lý Kinh Doanh & Ban Công Nghệ An Phước'],
    ['Đối tượng sử dụng', 'Cửa Hàng Trưởng (Store), Quản Lý Kinh Doanh (ASM), Admin / BGĐ'],
    ['Tính năng nổi bật', 'Traffic Realtime, P.QLQT Compliance, Dynamic HR, GIS Map, Excel 11 Sheets']
]
add_styled_table(doc, meta_headers, meta_rows, col_widths=[2.5, 4.0])

doc.add_page_break()

# ---------------------------------------------------------
# CHƯƠNG 1: TỔNG QUAN & ĐĂNG NHẬP PHÂN QUYỀN
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 1: TỔNG QUAN & ĐĂNG NHẬP PHÂN QUYỀN (RBAC)')

add_body_p(doc, 'Hệ thống An Phước Retail Commander Web Portal được xây dựng với cơ chế bảo mật cô lập dữ liệu theo vai trò (Strict Role-Based Scoping). Mọi yêu cầu truy vấn và lưu dữ liệu đều đi qua bộ lọc xác thực get_auth_scope() để đảm bảo quyền riêng tư và an toàn dữ liệu 100%.')

add_heading_2(doc, '1.1. Ma Trận Phân Quyền 3 Cấp')
auth_headers = ['Vai Trò Người Dùng', 'Phạm Vi Quyền Hạn (Scope)', 'Cơ Chế Đăng Nhập & Bảo Mật']
auth_rows = [
    ['Cửa Hàng Trưởng (Store)', 'Chỉ xem & nhập đúng 1 Cửa Hàng của mình', 'Chọn ASM ➔ Chọn CH ➔ Nhập PIN Store (4 số). Lock menu GIS Map.'],
    ['Quản Lý Kinh Doanh (ASM)', 'Xem & quản lý tất cả CH trong cụm phụ trách', 'Chọn ASM ➔ Dòng 🔑 [ASM] ➔ Nhập PIN ASM (9999). Xem Dashboard cụm & duyệt tờ trình.'],
    ['Admin & Ban Giám Đốc', 'Xem & quản lý TOÀN QUỐC (184 cửa hàng)', 'Chọn Admin ➔ Nhập PIN Admin/Master (8888). Xuất Excel Toàn Quốc & Đổi ASM động.']
]
add_styled_table(doc, auth_headers, auth_rows, col_widths=[2.0, 2.3, 2.2])

add_image_figure(doc, '_doc_assets/fig_auth_flow.png', 'Hình 1: Sơ đồ Kiến trúc Đăng nhập & Phân quyền Scoping (Strict Role Isolation)')

add_callout_box(doc, 'LƯU Ý QUAN TRỌNG VỀ ĐĂNG NHẬP CỬA HÀNG', 
                'Dữ liệu cửa hàng được cô lập hoàn toàn. Khi đăng nhập thành công, hệ thống lưu tự động thông tin vào localStorage. Nếu Cửa Hàng tải lại trang, hệ thống tự động phục hồi context mà không bị mất kết nối.', 
                bg_hex='EFF6FF', border_hex='2563EB', icon='🔐')

# ---------------------------------------------------------
# CHƯƠNG 2: NHẬP TRAFFIC & BÁO CÁO TUẦN (STORE)
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 2: HƯỚNG DẪN NHẬP TRAFFIC & BÁO CÁO TUẦN (DÀNH CHO STORE)')

add_heading_2(doc, '2.1. Nhập Số Liệu Traffic Hàng Ngày & Công Thức CR % (Tab 2)')
add_body_p(doc, 'Cửa hàng thực hiện nhập lượt khách và hóa đơn bán lẻ hàng ngày để dữ liệu báo cáo tuần tự động chính xác:')
add_bullet_p(doc, 'Khách (Traffic): ', 'Tổng số lượt khách ghé cửa hàng thực tế trong ngày (theo camera đếm lượt khách).')
add_bullet_p(doc, 'Bill (POS): ', 'Số lượng hóa đơn bán lẻ phát sinh thực tế tại quầy.')
add_bullet_p(doc, 'OL C.ty & OL CH: ', 'Phân tách số lượng hóa đơn Online công ty đẩy về đóng gói và đơn online do CH tự bán.')

add_body_p(doc, 'Công thức tính Tỷ lệ chuyển đổi (CR %) thực tế tại quầy:')
add_callout_box(doc, 'CÔNG THỨC CR % CHUẨN', 'Tỷ lệ CR (%) = [(Tổng Bill POS - Bill Online C.ty) / Lượt Khách Traffic] x 100%', bg_hex='ECFDF5', border_hex='059669', icon='📐')

add_heading_2(doc, '2.2. Quy Trình Lưu Traffic Mới (In-Page Banner & Soft Confirmation)')
add_body_p(doc, 'Khi số bill vượt quá lượt khách (CR > 100%) do camera đếm thiếu hoặc 1 khách mua nhiều bill, hệ thống gỡ bỏ hoàn toàn việc chặn cứng alert:')
add_bullet_p(doc, 'Banner Cảnh Báo In-Page: ', 'Hiển thị banner màu hổ phách/đỏ tổng hợp danh sách các ngày bất thường trực tiếp trên màn hình, không bị trình duyệt chặn alert popup.')
add_bullet_p(doc, 'Xác Nhận Lưu 100%: ', 'Cửa hàng bấm [Xác Nhận Lưu], số liệu lập tức được upsert thành công 100% lên Database Cloud Neon.')

add_image_figure(doc, '_doc_assets/fig_traffic_cr.png', 'Hình 2: Quy trình Nhập Traffic & Banner Cảnh Báo CR > 100% Không Bị Chặn Cứng')

add_heading_2(doc, '2.3. Nộp Báo Cáo Vận Hành Tuần (Tab 1)')
add_body_p(doc, 'Báo cáo tuần chốt số liệu vào Thứ Sáu hàng tuần gồm các mục chính:')
add_bullet_p(doc, 'Mục 3.1 & 3.2 (Hợp Đồng B2B): ', 'Điền hợp đồng phát sinh và hợp đồng cùng kỳ năm ngoái chưa ký lại (bắt buộc có Tên Khách Hàng / Doanh nghiệp).')
add_bullet_p(doc, 'Mục 4.5 (Yêu Cầu Hỗ Trợ 8 Cột): ', 'Ghi nhận sự cố kỹ thuật/POS/thiết bị. Hệ thống tự động đồng bộ 8 cột tương tác và gửi thông báo trực tiếp đến ASM.')

# ---------------------------------------------------------
# CHƯƠNG 3: MODULE P.QLQT SỰ VỤ TUÂN THỦ (TAB 3)
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 3: QUẢN LÝ SỰ VỤ TUÂN THỦ P.QLQT & TỜ TRÌNH (TAB 3)')

add_body_p(doc, 'Hệ thống tích hợp mô-đun kiểm soát tuân thủ và giải trình minh bạch 3 bên (Cửa Hàng ↔ ASM ↔ P.QLQT):')

add_heading_2(doc, '3.1. Cửa Hàng Giải Trình & Đính Kèm Tờ Trình Base64 DB Cloud')
add_body_p(doc, 'Cửa hàng xem vi phạm, mức phạt chức danh (CHT/CHP/NVBH), số tiền trừ thưởng và nộp giải trình:')
add_bullet_p(doc, 'Bản Giải Trình: ', 'Nhập nguyên nhân khách quan/chủ quan phát sinh vi phạm.')
add_bullet_p(doc, 'Tờ Trình Đính Kèm: ', 'Tải lên file Tờ trình (PDF, DOCX, PNG, JPG). File được mã hóa nhị phân Base64 lưu trực tiếp vào bảng tb_compliance_attachments trong Database Cloud Neon. Đảm bảo 100% không bao giờ bị mất file hay lỗi 404 khi máy chủ Render restart.')

add_image_figure(doc, '_doc_assets/fig_compliance_attachment.png', 'Hình 3: Kiến trúc Sự Vụ Tuân Thủ P.QLQT & Lưu Trữ Tờ Trình Base64 DB Cloud')

add_heading_2(doc, '3.2. ASM Đánh Giá & Chuyển Màu Trạng Thái Linh Hoạt')
add_body_p(doc, 'ASM mở Modal Đánh Giá Sự Vụ để xem giải trình và duyệt tờ trình:')
add_bullet_p(doc, 'Nút 👁️ Xem/Tải Tờ Trình: ', 'Xem trực tiếp hoặc tải file Tờ trình của Cửa Hàng ngay trên giao diện.')
add_bullet_p(doc, 'Đánh Giá Hoàn Tất (Emerald): ', 'Đồng ý với giải trình của Cửa Hàng. Nút bấm tự động chuyển sang màu xanh Emerald.')
add_bullet_p(doc, 'Đánh Giá Yêu Cầu Giải Trình Lại (Amber): ', 'Yêu cầu Cửa Hàng bổ sung hồ sơ. Nút bấm tự động chuyển sang màu Hổ Phách.')

# ---------------------------------------------------------
# CHƯƠNG 4: MODULE NHÂN SỰ NÂNG CẤP (TAB 4)
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 4: QUẢN LÝ ĐỊNH BIÊN & HỒ SƠ NHÂN SỰ NÂNG CẤP (TAB 4)')

add_heading_2(doc, '4.1. Bảng Hồ Sơ 6 Cột Cân Đối & Profile Ảnh Nén 1/4')
add_body_p(doc, 'Mô-đun nhân sự quản lý 1,089 nhân sự toàn hệ thống với các cải tiến vượt trội:')
add_bullet_p(doc, 'Sắp Xếp Ưu Tiên Chức Danh: ', 'Tự động đưa Chức danh chủ chốt lên đầu (CHT ➔ CHP ➔ NVBH ➔ Thử việc), sau đó sắp xếp theo Thâm niên công tác.')
add_bullet_p(doc, 'Profile Ảnh Chân Dung 1/4 Popup: ', 'Bấm vào tên nhân viên bất kỳ để mở Popup Hồ sơ với ảnh chân dung lớn chiếm 1/4 popup.')
add_bullet_p(doc, 'Tải Ảnh Nén Client-Side: ', 'Hỗ trợ chụp/tải ảnh trực tiếp từ điện thoại/máy tính có tự động nén dung lượng client-side.')
add_bullet_p(doc, 'Trường Ghi Chú Nâng Cao (notes): ', 'Lưu trữ chi tiết lịch sử điều chuyển cửa hàng, khen thưởng, kỷ luật, kỹ năng và đánh giá.')

add_image_figure(doc, '_doc_assets/fig_hr_architecture.png', 'Hình 4: Mô-đun Quản Lý Hồ Sơ Nhân Sự 6 Cột, Profile 1/4 & Loại Trừ Bảo Vệ')

add_heading_2(doc, '4.2. Quy Tắc Loại Trừ Bảo Vệ Khỏi Định Biên Headcount Quota')
add_callout_box(doc, 'QUY TẮC ĐỊNH BIÊN HEADCOUNT MỚI', 
                'Nhân sự có chức danh "Bảo vệ" được tự động loại trừ khỏi chỉ tiêu định biên bán hàng (Headcount quota) của Cửa Hàng, giúp phản ánh chính xác 100% lực lượng bán lẻ thực tế thiếu/đủ.', 
                bg_hex='FEF3C7', border_hex='D97706', icon='🛡️')

# ---------------------------------------------------------
# CHƯƠNG 5: CÔNG CỤ DÀNH CHO ASM & ADMIN
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 5: CÔNG CỤ DÀNH CHO ASM & ADMIN (DASHBOARD & EXCEL EXPORT)')

add_heading_2(doc, '5.1. Bảng Theo Dõi Dashboard ASM & Đổi ASM Quản Lý Động')
add_bullet_p(doc, 'Dashboard ASM: ', 'Theo dõi tiến độ nộp báo cáo tuần của từng cửa hàng trong cụm (Đã nộp - Xanh / Chưa nộp - Đỏ).')
add_bullet_p(doc, 'Đổi ASM Động (Admin Only): ', 'Admin có thể thay đổi ASM phụ trách cửa hàng trực tiếp trên giao diện, tự động đồng bộ lên Cloud Server qua API /api/admin/update_store_asm.')

add_heading_2(doc, '5.2. Xuất Báo Cáo Excel 11 Sheets Với Tên File Động')
add_body_p(doc, 'Bấm [📊 Xuất Báo Cáo Excel] để tải file báo cáo 11 Sheets chuẩn doanh nghiệp với màu sắc Navy sang trọng:')
add_bullet_p(doc, 'Tên File Xuất Theo Scope: ', 'Cửa hàng: BaoCao_RetailCommander_CH_126_3T2_2026-08-08.xlsx | ASM: BaoCao_RetailCommander_ASM_Khoi_2026-08-08.xlsx | Admin: BaoCao_RetailCommander_ToanQuoc_2026-08-08.xlsx')

# ---------------------------------------------------------
# CHƯƠNG 6: FAQ & TROUBLESHOOTING
# ---------------------------------------------------------
add_heading_1(doc, 'CHƯƠNG 6: HƯỚNG DẪN XỬ LÝ SỰ CỐ & CÂU HỎI THƯỜNG GẶP (FAQ)')

faq_headers = ['Tình Huống Bất Thường', 'Nguyên Nhân Kỹ Thuật', 'Cách Xử Lý Triệt Để']
faq_rows = [
    ['Bấm Lưu Traffic xuất hiện Banner CR > 100%', 'Số bill lớn hơn lượt khách (do đếm thiếu hoặc khách mua nhiều bill)', 'Xem danh sách ngày bất thường trên banner, bấm [Xác Nhận Lưu] để lưu 100% thành công.'],
    ['Bảo vệ có bị tính vào thiếu/đủ định biên không?', 'Quy tắc định biên headcount mới', 'Bảo vệ được tự động loại trừ khỏi định biên bán hàng để tính đúng lực lượng bán lẻ.'],
    ['Tải ảnh nhân viên từ điện thoại có bị nặng không?', 'Ảnh chụp từ camera có dung lượng lớn', 'Hệ thống tự động nén dung lượng client-side trước khi tải lên, đảm bảo siêu nhẹ & siêu nhanh.'],
    ['Tờ trình giải trình có bị mất khi server restart?', 'Lưu trữ dự trữ Database Cloud', 'File được mã hóa Base64 lưu trực tiếp vào Neon Postgres DB Cloud, bền vững 100%.'],
    ['Tài khoản Cửa Hàng không có nút Đánh Giá ASM?', 'Phân quyền bảo mật strict RBAC', 'Nút Đánh giá sự vụ chỉ dành riêng cho ASM & Admin. Cửa Hàng thực hiện Giải Trình & Đính Kèm Tờ Trình.']
]
add_styled_table(doc, faq_headers, faq_rows, col_widths=[2.2, 2.2, 2.3])

# Save
output_path1 = 'HUONG_DAN_SU_DUNG_WEB_PORTAL.docx'
output_path2 = r'C:\All_Report\8_RETAIL_COMMANDER\Retail_Commander_OS\10_PUBLICATION\HDSD_Retail_Commander_Web_Portal.docx'

doc.save(output_path1)

os.makedirs(os.path.dirname(output_path2), exist_ok=True)
doc.save(output_path2)

print('DOCX files successfully generated:')
print('1.', os.path.abspath(output_path1))
print('2.', os.path.abspath(output_path2))
